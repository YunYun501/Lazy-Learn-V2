import pytest
import tempfile
from pathlib import Path
from docx import Document
from app.services.docx_parser import DOCXParser


def create_test_docx(filepath: str, content: list[dict]):
    """Create a test DOCX with headings and paragraphs.
    content: list of {"type": "heading1"|"heading2"|"paragraph", "text": "..."}
    """
    doc = Document()
    for item in content:
        if item["type"] == "heading1":
            doc.add_heading(item["text"], level=1)
        elif item["type"] == "heading2":
            doc.add_heading(item["text"], level=2)
        else:
            doc.add_paragraph(item["text"])
    doc.save(filepath)


def test_docx_detects_headings():
    """Test that DOCX parsing detects heading levels."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name

    content = [
        {"type": "heading1", "text": "Chapter 1: Introduction"},
        {"type": "paragraph", "text": "This chapter introduces the Z-transform."},
        {"type": "heading2", "text": "1.1 Definition"},
        {"type": "paragraph", "text": "The Z-transform is defined as..."},
    ]
    create_test_docx(tmp_path, content)

    parser = DOCXParser()
    sections = parser.parse(tmp_path)

    assert len(sections) >= 2
    # First section should be under "Chapter 1: Introduction"
    assert any("Introduction" in s.section for s in sections)


def test_docx_preserves_text():
    """Test that paragraph text is preserved."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name

    content = [
        {"type": "heading1", "text": "Z-Transform"},
        {"type": "paragraph", "text": "The Z-transform converts discrete-time signals."},
    ]
    create_test_docx(tmp_path, content)

    parser = DOCXParser()
    sections = parser.parse(tmp_path)

    all_text = " ".join(s.text for s in sections)
    assert "Z-transform" in all_text or "discrete-time" in all_text


def test_docx_to_chapters_format():
    """Test that to_chapters returns correct format."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = f.name

    content = [
        {"type": "heading1", "text": "Introduction"},
        {"type": "paragraph", "text": "Content here."},
    ]
    create_test_docx(tmp_path, content)

    parser = DOCXParser()
    sections = parser.parse(tmp_path)
    chapters = parser.to_chapters(sections)

    assert len(chapters) >= 1
    assert "number" in chapters[0]
    assert "title" in chapters[0]
    assert "text" in chapters[0]
